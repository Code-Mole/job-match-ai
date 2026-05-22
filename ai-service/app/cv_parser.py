"""
CV Parser — extracts raw text and identifies skills from uploaded CV files.
Supports PDF (via pdfplumber + PyMuPDF fallback), DOCX, and TXT.
"""

import re
import os
import logging
from typing import Optional

# PDF extraction libraries — pdfplumber is more accurate, fitz is faster
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX extraction
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .cv_skill_extractor import extract_cv_profile, extract_skills_from_cv_text

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    Tries pdfplumber first (better layout handling), falls back to PyMuPDF.
    """
    text = ""

    # Method 1: pdfplumber — better for tables and columns
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF...")

    # Method 2: PyMuPDF fallback
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        except Exception as e:
            logger.error(f"PyMuPDF also failed: {e}")

    raise RuntimeError("No PDF library available. Install pdfplumber or pymupdf.")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file — reads all paragraphs and tables."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed.")

    doc = Document(file_path)
    parts = []

    # Main paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Text inside tables (many CVs use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_text_from_file(file_path: str) -> str:
    """
    Dispatch to the right extractor based on file extension.
    Returns raw text string.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in (".txt", ".text"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_skills_from_text(text: str) -> list:
    """Open CV-based skill extraction (all industries, not tech-only)."""
    return extract_skills_from_cv_text(text)


def extract_years_of_experience(text: str) -> int:
    """
    Try to infer total years of experience from CV text.
    Looks for patterns like "5 years of experience", "3+ years", etc.
    Returns the highest number found (as a conservative estimate).
    """
    patterns = [
        r'(\d+)\+?\s+years?\s+of\s+(?:professional\s+)?experience',
        r'(\d+)\+?\s+years?\s+(?:of\s+)?(?:work\s+)?experience',
        r'experience\s+of\s+(\d+)\+?\s+years?',
        r'(\d+)\+?\s+years?\s+(?:in|with|using)',
    ]
    years_found = []
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        years_found.extend(int(m) for m in matches)

    return max(years_found) if years_found else 0


def parse_cv(file_path: str) -> dict:
    """
    Main entry point — parse a CV file and return structured data.

    Returns:
        {
            "raw_text": str,
            "skills": list[str],
            "years_experience": int,
            "word_count": int,
            "char_count": int,
        }
    """
    try:
        raw_text = extract_text_from_file(file_path)
        profile = extract_cv_profile(raw_text)
        skills = profile["skills"]
        years_exp = extract_years_of_experience(raw_text)

        return {
            "success":          True,
            "raw_text":         raw_text,
            "skills":           skills,
            "roles":            profile.get("roles", []),
            "strengths":        profile.get("strengths", {}),
            "years_experience": years_exp,
            "word_count":       len(raw_text.split()),
            "char_count":       len(raw_text),
        }

    except Exception as e:
        logger.error(f"CV parsing failed for {file_path}: {e}")
        return {
            "success": False,
            "error":   str(e),
            "raw_text": "",
            "skills":   [],
            "years_experience": 0,
        }